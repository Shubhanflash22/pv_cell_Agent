#!/usr/bin/env python3
"""Generate the methodology report as a .docx file."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_paragraph(doc, text):
    return doc.add_paragraph(text)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
    for row_idx, row in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, val in enumerate(row):
            row_cells[col_idx].text = str(val)
    return table

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading("SolarInvest: Methodology Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("AI-Powered Residential PV Sizing")
    doc.add_paragraph()
    
    # Section 1
    add_heading(doc, "1. Problem Formulation and Core Formulation", 1)
    add_heading(doc, "1.1 Problem Statement", 2)
    add_paragraph(doc, 
        "The system addresses the residential solar sizing problem: given a homeowner's location, "
        "budget, roof dimensions, and household profile, recommend a PV system (panel count, brand, "
        "and optional battery) that balances technical feasibility, budget constraints, and financial returns.")
    add_paragraph(doc, 
        "The decision variables are: N (panel count), Panel brand (from a catalog of 9 manufacturers), "
        "and Battery (add / evaluate later / PV-only). Subject to: Roof capacity (N ≤ N_roof), "
        "Budget (CAPEX ≤ budget_usd), and User brand preference (if specified).")
    
    add_heading(doc, "1.2 Core Formulation", 2)
    add_paragraph(doc, 
        "The model produces two scenarios. Optimal scenario (technically best): "
        "N_opt = min(N_100%, N_roof) where N_100% is panels for 100% annual consumption offset. "
        "Recommended scenario (budget-aware): N_rec = min(N_70%, N_budget, N_roof) where N_70% targets "
        "70% offset and N_budget is max panels affordable at the given budget.")
    add_paragraph(doc, 
        "The financial objective is 10-year NPV: NPV = -Net_CAPEX + Σ(Savings_y / (1+r)^y) with r = 7% "
        "discount rate. Net CAPEX is post-ITC (30% federal credit).")
    
    # Section 2
    add_heading(doc, "2. Input Data: Extraction, Processing, and Modelling", 1)
    add_heading(doc, "2.1 Weather Data", 2)
    add_paragraph(doc, 
        "Source: Open-Meteo Archive API. Variables: daily (temperature max/min, shortwave radiation sum), "
        "hourly (cloud cover, shortwave radiation). Processing: weather_fetcher.py fetches 1–5 years of "
        "history, aggregates to weekly (weekly_max_temperature, weekly_avg_irradiance, weekly_avg_cloud_cover). "
        "fetch_irradiance_annual() derives mean annual GHI (kWh/m²/yr). Fallback: 2080 kWh/m²/yr.")
    
    add_heading(doc, "2.2 Household Usage Data", 2)
    add_paragraph(doc, 
        "Source: EIA San Diego regional load (San_Diego_Load_EIA_Fixed.csv), hourly MW. Processing: "
        "avg_kw = MW_Load × 1000 / SDGE_TOTAL_CUSTOMERS (1,040,149 customers). Nine variability factors "
        "applied: F1 longitude (coastal→inland 0.85×–1.25×), F2 latitude (0.90×–1.10×), F3 elevation, "
        "F4 household size, F5 urban density (0.7×–1.3×), F6 economic proxy, F7 solar adoption, F8 EV "
        "charging (7.2 kW per EV), F9 multi-generational. Deterministic via SHA-256 hash of (lat, lon).")
    
    add_heading(doc, "2.3 Tariff Modelling", 2)
    add_paragraph(doc, 
        "Source: SDG&E TOU rate CSVs (tou_dr_daily_2021_2025.csv). Periods: On-peak 16:00–21:00, "
        "Super-off-peak 00:00–06:00 and Mar–Apr 10:00–14:00, Off-peak otherwise. build_hourly_tariffs() "
        "maps each of 8760 hours to the correct rate.")
    
    # Section 3
    add_heading(doc, "3. Feature Engineering", 1)
    add_paragraph(doc, 
        "feature_engineering.py derives ~70 features in seven categories: (1) Electricity: load "
        "distribution, seasonality, growth, peak load. (2) Weather/Solar: irradiance stats, PV efficiency, "
        "consumption–irradiance correlation. (3) Household: kWh per occupant, cost, annual spend. "
        "(4) Cross-dataset: panels for 100%/70%/50% offset, break-even, NPV, IRR, ROI, nighttime load "
        "ratio. (5) Risk: price and irradiance sensitivity. (6) EV & Budget: EV charging load, budget "
        "analysis. (7) format_for_llm() produces a structured text block for the prompt.")
    
    # Section 4
    add_heading(doc, "4. Mathematical Optimisation", 1)
    add_heading(doc, "4.1 Constants", 2)
    add_table(doc, 
        ["Constant", "Value", "Source"],
        [
            ["G_ref", "1000 W/m²", "STC"],
            ["PR (performance ratio)", "0.80", "pv_tools.py"],
            ["Federal ITC", "30%", "config.yaml"],
            ["Utility escalation", "6%/yr", "pv_tools.py"],
            ["Discount rate", "7%", "pv_tools.py"],
            ["O&M", "$0.005/W/yr", "pv_tools.py"],
            ["NEM export credit", "$0.10/kWh", "pv_tools.py"],
            ["Inverter replacement", "$2000 at year 10", "pv_tools.py"],
            ["SDG&E daily fee", "$0.345/day", "pv_tools.py"],
            ["EV charger power", "7.2 kW", "pv_tools.py"],
            ["EV daily energy", "14 kWh/EV", "pv_tools.py"],
        ])
    
    add_heading(doc, "4.2 Core Functions", 2)
    add_paragraph(doc, 
        "Irradiance shape: G(t)/G_ref = max(0, sin(π(h - h_sunrise)/(h_sunset - h_sunrise))) with "
        "h_sunset = 18 + 2·sin(2π(doy-80)/365). PV output: P_array(t) = N × P_panel × shape(t) × PR, "
        "normalised so annual sum = N × P_panel × GHI × PR. Dispatch: surplus → charge battery then "
        "export; deficit → discharge battery then import. Economics: Net_CAPEX = (PV + Battery) × 1.10 × "
        "0.70; Savings_y = Trad_bill_y - Solar_bill_y - O&M - Inverter; with degradation (1-d)^(y-1) "
        "and tariff escalation (1+0.06)^(y-1).")
    
    add_heading(doc, "4.3 Hardware Selection", 2)
    add_paragraph(doc, 
        "Panel: select_panel(brand) — auto = best efficiency/cost; user = exact manufacturer match. "
        "Battery: select_battery(required_kwh) — min cost/kWh when required_kwh > 0.5; else None. "
        "Battery need: battery_kwh = nighttime_fraction × annual_kwh/365 when nighttime_fraction > 0.30.")
    
    # Section 5
    add_heading(doc, "5. Stochastic Scenario", 1)
    add_paragraph(doc, 
        "Determinism: household load is seeded by hashlib.sha256(lat_lon). Same location always yields "
        "identical profile. Stochastic elements: location (coastal vs inland, urban vs suburban), solar "
        "adoption (5–35% prob), EV adoption (8–30% prob), multi-generational (10–25% prob), household "
        "characteristics (normal draws), ±3% hourly noise. All pseudo-random; fixed per (lat, lon).")
    
    # Section 6
    add_heading(doc, "6. Hierarchical Prompt Structure", 1)
    add_paragraph(doc, 
        "prompt_builder.py assembles the prompt in order: (1) Feature summary (~2,500 chars), "
        "(2) Equipment catalog (9 panels, 5 batteries, constants, EV assumptions), (3) User inputs, "
        "(4) Pre-computed tool results (panel, brand, battery, load, tariff, roof layout, sizing, "
        "recommended/optimal scenarios, battery analysis), (5) Hard rules (11 anti-hallucination "
        "constraints), (6) Decision policy (copy instructions), (7) JSON schema, (8) Task. Truncation: "
        "if > 24,000 chars, features truncated first; tool results, rules, schema, task protected.")
    
    # Section 7
    add_heading(doc, "7. Constraints Enforced", 1)
    add_paragraph(doc, 
        "Hard rules: (1) Numeric integrity — copy from TOOL RESULTS > FEATURES > CATALOG > USER INPUTS; "
        "no own arithmetic. (2) Scenario structure — exactly two scenarios. (3) Evidence — 5–12 entries. "
        "(4) Brand — use user preference when specified. (5) Roof — N ≤ max_panels_by_roof. (6) Tariff — "
        "use rates from TOOL RESULTS. (7) Output — valid JSON only. (8) Self-check — panels×Wp/1000 = "
        "kw_dc. (9) Battery — copy from BATTERY ANALYSIS. (10) Panel brand — copy from BRAND SELECTION. "
        "Decision policy: N_opt = min(N_100, N_roof), N_rec = min(N_70, N_budget, N_roof), "
        "budget_binding = True if N_budget < N_70.")
    
    # Section 8
    add_heading(doc, "8. Equipment Catalogs", 1)
    add_heading(doc, "8.1 Solar Panels (9)", 2)
    add_table(doc, 
        ["Manufacturer", "Model", "Wp", "Efficiency", "$/Wp", "Degrad/yr"],
        [
            ["REC Group", "Alpha Pure", "405", "22.6%", "2.85", "0.005"],
            ["JA Solar", "DeepBlue", "395", "21.5%", "2.90", "0.006"],
            ["Trina Solar", "Vertex S", "400", "21.8%", "2.90", "0.006"],
            ["Canadian Solar", "TOPHiKu7", "420", "22.5%", "3.10", "0.005"],
            ["Silfab Solar", "Prime", "410", "22.1%", "3.15", "0.005"],
            ["Jinko Solar", "Tiger Neo", "440", "23.8%", "3.20", "0.005"],
            ["LONGi Solar", "Hi-MO 6", "435", "23.3%", "3.35", "0.005"],
            ["Maxeon Solar", "Maxeon 7", "430", "22.8%", "3.50", "0.004"],
            ["Aiko Solar", "Neostar 2P", "460", "24.3%", "3.75", "0.004"],
        ])
    
    add_heading(doc, "8.2 Batteries (5)", 2)
    add_table(doc, 
        ["Manufacturer", "Model", "kWh", "Cost", "RTE"],
        [
            ["Tesla", "Powerwall 3", "13.5", "$11,500", "97.5%"],
            ["Enphase", "IQ Battery 5P", "5.0", "$6,000", "96.0%"],
            ["Generac", "PWRcell M6", "9.0", "$10,000", "96.5%"],
            ["SolarEdge", "Home Battery 48V", "9.7", "$9,500", "94.5%"],
            ["Panasonic", "EverVolt H Series", "17.1", "$15,000", "97.0%"],
        ])
    
    # Section 9
    add_heading(doc, "9. Two Challenges and Solutions", 1)
    add_heading(doc, "Challenge 1: LLM Hallucination of Financial Numbers", 2)
    add_paragraph(doc, 
        "Problem: LLMs invent or round numbers. Solution: Pre-compute all numbers in pv_tools.run_all_tools() "
        "and pass them to the LLM. The LLM's job is to copy and format, not calculate. Hard rules and "
        "decision policy explicitly forbid arithmetic and require copying from TOOL RESULTS.")
    
    add_heading(doc, "Challenge 2: Household Load Without Real Data", 2)
    add_paragraph(doc, 
        "Problem: No per-household metered data for arbitrary locations. Solution: Use EIA regional load "
        "as base, downscale by SDGE customer count, apply nine location-based factors. SHA-256 seeded RNG "
        "ensures same (lat, lon) + household params always yields identical profile. Deterministic, "
        "reproducible, location-specific household profiles.")
    
    # Section 10
    add_heading(doc, "10. Core Objects and Data Flow", 1)
    add_paragraph(doc, 
        "Core objects: user_inputs (lat, lon, num_evs, num_people, num_daytime_occupants, budget_usd, "
        "roof_length_m, roof_breadth_m, rate_plan, panel_brand), tool_results (panel, brand, battery, "
        "load, tariff, roof layout, sizing, recommended/optimal scenarios, battery analysis), "
        "recommendation (optimal, recommended, battery_recommendation, panel_brand_recommendation, evidence).")
    add_paragraph(doc, 
        "Flow: User inputs → extraction → features → PV tools → prompt → LLM → validation → report.")
    
    doc.add_paragraph()
    add_paragraph(doc, 
        "— This report is based on the SolarInvest codebase: config.yaml, pv_tools.py, "
        "feature_engineering.py, prompt_builder.py, data_extractor.py, weather_fetcher.py, "
        "household_generator.py, and the README documentation.")
    
    out_path = "outputs/SolarInvest_Methodology_Report.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == "__main__":
    main()
